"""DOCX parser tests — Technical_Architecture_v2.md SS3, Phase 2 spec SS5."""

from __future__ import annotations

from io import BytesIO

import docx as docxlib

from app.models.enums import ErrorCode
from app.services.parsing.docx_parser import parse_docx
from app.services.parsing.models import ParsedDocument, ParseResult, ParseStatus
from tests.fixtures.synthetic_documents import (
    SYNTHETIC_CLAUSES,
    build_corrupted_docx,
    build_docx,
    build_empty_docx,
)

_DEFAULTS: dict[str, int] = {"max_items": 5000, "min_text_chars": 50}


def _require_document(result: ParseResult) -> ParsedDocument:
    assert result.document is not None
    return result.document


def test_normal_docx_parses_successfully():
    result = parse_docx(build_docx(), **_DEFAULTS)
    assert result.status == ParseStatus.SUCCESS
    assert result.is_success
    assert result.error_code is None
    document = _require_document(result)
    assert document.total_text_length > 0
    assert document.page_count is None  # DOCX has no native page concept


def test_paragraph_order_preserved():
    document = _require_document(parse_docx(build_docx(), **_DEFAULTS))
    text_blocks = [b for b in document.blocks if not b.is_table_cell]
    assert [b.order for b in text_blocks] == sorted(b.order for b in text_blocks)
    assert text_blocks[0].text == SYNTHETIC_CLAUSES[0][1]
    assert text_blocks[1].text == SYNTHETIC_CLAUSES[1][1]


def test_heading_style_preserved():
    document = _require_document(parse_docx(build_docx(), **_DEFAULTS))
    headings = [b for b in document.blocks if b.is_heading]
    assert len(headings) == 3  # SYNTHETIC_CLAUSES has 3 "Heading 1" entries
    assert all(b.style == "Heading 1" for b in headings)


def test_table_contents_extracted_without_crashing():
    document = _require_document(parse_docx(build_docx(include_table=True), **_DEFAULTS))
    table_cells = [b for b in document.blocks if b.is_table_cell]
    cell_texts = {b.text for b in table_cells}
    assert "Fee Type" in cell_texts
    assert "$25 synthetic flat fee" in cell_texts


def test_document_without_table_has_no_table_cell_blocks():
    document = _require_document(parse_docx(build_docx(include_table=False), **_DEFAULTS))
    assert not any(b.is_table_cell for b in document.blocks)


def test_list_style_paragraph_flagged_as_list_item():
    source = docxlib.Document()
    source.add_paragraph("First synthetic list item about a fee.", style="List Bullet")
    buffer = BytesIO()
    source.save(buffer)

    parsed = _require_document(parse_docx(buffer.getvalue(), max_items=5000, min_text_chars=1))
    assert any(b.is_list_item for b in parsed.blocks)


def test_bold_run_flagged():
    source = docxlib.Document()
    paragraph = source.add_paragraph("Synthetic bold clause text about default.")
    paragraph.runs[0].bold = True
    buffer = BytesIO()
    source.save(buffer)

    parsed = _require_document(parse_docx(buffer.getvalue(), max_items=5000, min_text_chars=1))
    assert any(b.is_bold for b in parsed.blocks)


def test_empty_docx_is_low_text_content():
    result = parse_docx(build_empty_docx(), **_DEFAULTS)
    assert not result.is_success
    assert result.status == ParseStatus.EMPTY
    assert result.error_code == ErrorCode.LOW_TEXT_CONTENT


def test_near_empty_docx_is_low_text_content():
    document = docxlib.Document()
    document.add_paragraph("hi")
    buffer = BytesIO()
    document.save(buffer)

    result = parse_docx(buffer.getvalue(), max_items=5000, min_text_chars=50)
    assert not result.is_success
    assert result.status == ParseStatus.LOW_TEXT
    assert result.error_code == ErrorCode.LOW_TEXT_CONTENT


def test_corrupted_docx_returns_corrupted_status_not_exception():
    result = parse_docx(build_corrupted_docx(), **_DEFAULTS)
    assert result.status == ParseStatus.CORRUPTED
    assert result.error_code == ErrorCode.CORRUPTED_FILE
    assert "Traceback" not in (result.detail or "")


def test_malformed_package_returns_corrupted_not_exception():
    # Not a zip at all — a different failure surface than build_corrupted_docx().
    result = parse_docx(b"this is not a zip file at all", **_DEFAULTS)
    assert result.status == ParseStatus.CORRUPTED
    assert result.error_code == ErrorCode.CORRUPTED_FILE


def test_too_many_items_rejected():
    items = [("Normal", f"Synthetic paragraph number {i} about a fee.") for i in range(10)]
    result = parse_docx(build_docx(items=items, include_table=False), max_items=3, min_text_chars=1)
    assert result.status == ParseStatus.TOO_MANY_PAGES
    assert result.error_code == ErrorCode.FILE_TOO_LARGE

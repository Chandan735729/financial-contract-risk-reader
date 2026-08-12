"""Synthetic PDF/DOCX byte generators for Phase 2 tests.

Every document is built programmatically at test time — nothing binary is
committed to the repository, and every string here is clearly fictional
contract-style text written for this test suite only (Phase 2 spec SS12).
"""

from __future__ import annotations

import zipfile
from io import BytesIO

import docx
import pymupdf

SYNTHETIC_CLAUSES = [
    ("Heading 1", "1. Prepayment"),
    (
        "Normal",
        "Borrower may prepay this synthetic loan of $1,000 subject to a 2% "
        "prepayment penalty fee if repaid within 12 months of origination.",
    ),
    ("Heading 1", "2. Default and Acceleration"),
    (
        "Normal",
        "A default under this synthetic agreement triggers immediate "
        "acceleration of the outstanding balance at a rate of 8.5% per annum.",
    ),
    ("Heading 1", "3. Auto-Renewal"),
    (
        "Normal",
        "This synthetic policy automatically renews for a period of 12 months "
        "unless the policyholder provides 30 days written notice.",
    ),
]


def build_pdf(page_texts: list[str] | None = None) -> bytes:
    if page_texts is None:
        page_texts = [f"{heading}\n{body}" for heading, body in SYNTHETIC_CLAUSES]

    doc = pymupdf.open()
    for text in page_texts:
        page = doc.new_page()
        page.insert_text((72, 72), text, fontsize=11)
    data = doc.tobytes()
    doc.close()
    return data


def build_empty_pdf() -> bytes:
    # pymupdf refuses to save a literal 0-page document ("cannot save with
    # zero pages") — a single blank page with no text/images is the
    # practical, constructible equivalent of an "empty PDF" upload and
    # exercises the same LOW_TEXT_CONTENT path.
    doc = pymupdf.open()
    doc.new_page()
    data = doc.tobytes()
    doc.close()
    return data


def build_scanned_pdf(page_count: int = 2) -> bytes:
    """Structurally valid, multi-page, contains an embedded image and no
    extractable text — simulates a scanned/image-only document."""
    doc = pymupdf.open()
    pixmap = pymupdf.Pixmap(pymupdf.csRGB, (0, 0, 10, 10), False)
    pixmap.set_rect(pixmap.irect, (200, 200, 200))
    image_bytes = pixmap.tobytes("png")
    for _ in range(page_count):
        page = doc.new_page()
        page.insert_image(pymupdf.Rect(0, 0, 200, 200), stream=image_bytes)
    data = doc.tobytes()
    doc.close()
    return data


def build_password_protected_pdf() -> bytes:
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Synthetic confidential contract text.", fontsize=11)
    data = doc.tobytes(
        encryption=pymupdf.PDF_ENCRYPT_AES_256,  # type: ignore[attr-defined]  # real at runtime; pymupdf's stubs are incomplete
        owner_pw="synthetic-owner-pw",
        user_pw="synthetic-user-pw",
    )
    doc.close()
    return data


def build_corrupted_pdf() -> bytes:
    # Valid PDF magic bytes (passes content-sniffing) but a structurally
    # invalid body — distinct from build_random_binary(), which fails
    # sniffing itself.
    return b"%PDF-1.7\n%\xe2\xe3\xcf\xd3\nthis is not a real pdf body, just garbage\n%%EOF"


def build_docx(items: list[tuple[str, str]] | None = None, include_table: bool = True) -> bytes:
    """`items`: list of (style_name, text) tuples; style "Normal" means no
    explicit style override."""
    if items is None:
        items = SYNTHETIC_CLAUSES

    document = docx.Document()
    for style, text in items:
        document.add_paragraph(text, style=None if style == "Normal" else style)

    if include_table:
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Fee Type"
        table.cell(0, 1).text = "Amount"
        table.cell(1, 0).text = "Late Payment Fee"
        table.cell(1, 1).text = "$25 synthetic flat fee"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_empty_docx() -> bytes:
    buffer = BytesIO()
    docx.Document().save(buffer)
    return buffer.getvalue()


def build_corrupted_docx() -> bytes:
    """Structurally docx-shaped (passes content-sniffing — `[Content_Types].xml`
    and `word/document.xml` entries are present) but `word/document.xml` is
    garbled, so python-docx fails while parsing it. Distinct from
    build_random_binary(), which fails sniffing itself."""
    good = build_docx(items=[("Normal", "placeholder")], include_table=False)
    src = zipfile.ZipFile(BytesIO(good))
    out_buffer = BytesIO()
    with zipfile.ZipFile(out_buffer, "w") as out:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename == "word/document.xml":
                data = b"<not-valid-xml-at-all"
            out.writestr(item, data)
    return out_buffer.getvalue()


def build_random_binary(size: int = 256) -> bytes:
    # Fixed, deterministic "random" bytes — no PDF/ZIP magic anywhere in it.
    return bytes((i * 37 + 11) % 256 for i in range(size))

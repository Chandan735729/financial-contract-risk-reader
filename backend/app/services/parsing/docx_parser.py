"""DOCX parsing — Technical_Architecture_v2.md SS3 (python-docx).

Parses from an in-memory byte string only. Never trusts the `.docx`
extension as proof of validity — `docx.Document()` is given raw bytes and
whatever it raises (corrupted zip, missing OPC parts, not a Word package at
all) is caught and mapped to `ErrorCode.CORRUPTED_FILE`, never re-raised.
"""

from __future__ import annotations

from io import BytesIO

import docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.services.parsing.models import DocumentTextBlock, ParsedDocument, ParseResult, ParseStatus


def parse_docx(data: bytes, *, max_items: int, min_text_chars: int) -> ParseResult:
    try:
        document = docx.Document(BytesIO(data))
    except Exception:
        return ParseResult(status=ParseStatus.CORRUPTED, detail="failed to open as DOCX")

    try:
        return _parse_opened_docx(document, max_items=max_items, min_text_chars=min_text_chars)
    except Exception:
        # A malformed-but-openable package (e.g. a table cell referencing a
        # missing style) must not crash the request — treat as corrupted
        # rather than letting a parser internal escape to the client.
        return ParseResult(status=ParseStatus.CORRUPTED, detail="failed while reading DOCX content")


def _iter_block_items(document: DocxDocument):
    """Yield paragraphs and tables in document order (python-docx has no
    built-in for this — `document.paragraphs`/`document.tables` are each
    flat and lose interleaving, which matters for preserving reading order
    for Phase 3 segmentation)."""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _is_list_item(paragraph: Paragraph) -> bool:
    # Direct paragraph-level numbering (explicit numPr) *or* a built-in
    # "List..." style (e.g. "List Bullet", "List Number") — Word's built-in
    # list styles carry numbering via the style definition, not a numPr on
    # every paragraph, so the style name is the more reliable signal in
    # practice for documents built with those styles.
    if paragraph._p.xpath("./w:pPr/w:numPr"):
        return True
    style_name = paragraph.style.name if paragraph.style else ""
    return style_name.lower().startswith("list")


def _is_bold(paragraph: Paragraph) -> bool:
    return any(run.bold for run in paragraph.runs if run.bold is not None)


def _is_heading(paragraph: Paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style else ""
    return style_name.lower().startswith("heading") or style_name.lower() == "title"


def _iter_table_cells(table: Table):
    seen_cell_ids: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            # Merged cells appear multiple times in `row.cells` (python-docx
            # repeats the same underlying cell for each grid position it
            # spans) — de-dupe by identity so merged text isn't duplicated.
            if id(cell) in seen_cell_ids:
                continue
            seen_cell_ids.add(id(cell))
            yield cell


def _parse_opened_docx(document: DocxDocument, *, max_items: int, min_text_chars: int) -> ParseResult:
    blocks: list[DocumentTextBlock] = []
    order = 0
    char_offset = 0
    item_count = 0

    for item in _iter_block_items(document):
        item_count += 1
        if isinstance(item, Paragraph):
            text = item.text
            if text.strip():
                start = char_offset
                char_offset += len(text)
                blocks.append(
                    DocumentTextBlock(
                        text=text,
                        order=order,
                        source_type="docx",
                        start_char=start,
                        end_char=char_offset,
                        style=item.style.name if item.style else None,
                        is_heading=_is_heading(item),
                        is_list_item=_is_list_item(item),
                        is_bold=_is_bold(item),
                    )
                )
                order += 1
        elif isinstance(item, Table):
            for cell in _iter_table_cells(item):
                item_count += 1
                cell_text = _cell_text(cell)
                if cell_text.strip():
                    start = char_offset
                    char_offset += len(cell_text)
                    blocks.append(
                        DocumentTextBlock(
                            text=cell_text,
                            order=order,
                            source_type="docx",
                            start_char=start,
                            end_char=char_offset,
                            style="TableCell",
                            is_table_cell=True,
                        )
                    )
                    order += 1

        if item_count > max_items:
            return ParseResult(
                status=ParseStatus.TOO_MANY_PAGES,
                detail=f"exceeds max_items={max_items}",
            )

    total_chars = sum(len(b.text) for b in blocks)
    if total_chars == 0:
        return ParseResult(status=ParseStatus.EMPTY, detail="0 extracted chars")
    if total_chars < min_text_chars:
        return ParseResult(status=ParseStatus.LOW_TEXT, detail=f"{total_chars} extracted chars")

    return ParseResult(
        status=ParseStatus.SUCCESS,
        document=ParsedDocument(source_type="docx", blocks=tuple(blocks), page_count=None),
        detail=f"{total_chars} extracted chars",
    )


def _cell_text(cell: _Cell) -> str:
    # `_Cell.text` already joins the cell's paragraphs; guarded separately
    # so one malformed cell can't take down the whole table.
    try:
        return cell.text
    except Exception:
        return ""
